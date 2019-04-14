from bs4 import BeautifulSoup
import requests
import csv
from docx import Document
from docx.shared import Inches
from datetime import datetime
from datetime import timedelta

name = input("Please Name the Document ")
document = Document()
document.add_heading('Sözcü', 0)

source = requests.get('http://www.sozcu.com.tr/').text

soup = BeautifulSoup(source, 'lxml')
mainSwiper = soup.find('div',class_='swiper-wrapper clearfix')
swiper = mainSwiper.find_all('div', class_='swiper-slide')

i=0
List = []
Thumbnails = []
for news in swiper:
    List.append(news.a.get('href'))
    Thumbnails.append(news.img.get('src'))
    i = i + 1



for links in List:
    counter = 0
    inner_source = requests.get(links).text
    inner_soup = BeautifulSoup(inner_source,'lxml')
    article = inner_soup.find('div', class_='content-element')
    if article != None:
        paragraphs = article.find_all('p')
        print(paragraphs)
        for p in paragraphs:
            if(counter!=0):
                document.add_paragraph(p.text)
            else:
                document.add_paragraph(p, style='IntenseQuote')

            counter = counter+1

document.save(f'{name}.docx')